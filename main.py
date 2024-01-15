import json
from dataclasses import dataclass, asdict
from pathlib import Path
import re

import requests
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.shared import Pt
from tqdm import tqdm
from docx.enum.section import WD_SECTION_START

MIN_HSK_LEVEL = 4
MAX_WORDS_PER_COLUMN = 17

BASE_URL = (
    "https://mandarinbean.com/all-lessons/?jsf=epro-posts&tax=post_tag:12%2C9%2C20"
)
METADATA_JSON = Path("article_metadata.json")
FINAL_TEXT_JSON = Path("final_text.json")


HEADERS = {
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Sec-Fetch-Site": "same-origin",
    # 'Cookie': '_ga=GA1.1.119314374.1689522493; _ga_JMQ35CXBDN=GS1.1.1689522492.1.1.1689523449.0.0.0; _gid=GA1.2.1441052773.1689522493; _gat_UA-117095528-1=1',
    "Sec-Fetch-Dest": "document",
    "Accept-Language": "en-US,en;q=0.9",
    "Sec-Fetch-Mode": "navigate",
    "Host": "mandarinbean.com",
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.5.2 Safari/605.1.15",
    "Referer": "https://mandarinbean.com/all-lessons/",
    # 'Accept-Encoding': 'gzip, deflate, br',
    "Connection": "keep-alive",
}


@dataclass
class ArticleMetadata:
    """Metadata for an article. This includes the title, chinese title, url, tags, and HSK level"""

    title: str
    chinese_title: str
    url: str
    tags: str
    hsk_tag: str


@dataclass
class ArticleTextCollection(ArticleMetadata):
    """The main text and word list for an article, along with the metadata"""

    main_text: str
    word_list: str

    @property
    def sanitized_word_list(self) -> set[str]:
        words = set(self.word_list.split("\n"))
        sanitized = {w.lstrip("\n").rstrip('"n').strip() for w in words if w}

        # remove any (HSK...) tags from the words
        return {re.sub(r"\(HSK\d+\)", "", w) for w in sanitized}


class MainArticleFetcher:
    """Converts a Mandarin Bean article URL to a text file with the main text and a word list. The words aren't just
    scraped directly from the page, but instead are extracted from the tooltip text. This is because the tooltip text
    contains the HSK level of the word, to build a word list.
    """

    def __call__(self, url: str) -> tuple[str, str]:
        response = requests.get(url, headers=HEADERS)
        page_soup = BeautifulSoup(response.text, "html.parser")

        main_text = self._soup_to_main_text(page_soup)
        word_list = self._soup_to_word_list(page_soup)

        return main_text, word_list

    def _soup_to_main_text(self, soup: BeautifulSoup) -> str:
        paragraphs = soup.find_all("p")
        return "\n".join(_p for p in paragraphs if (_p := self._paragraph_to_text(p)))

    @staticmethod
    def _paragraph_to_text(paragraph: Tag) -> str:
        """Converts a paragraph tag to text. This is done by extracting the punctuation and tr characters from the
        tooltip text, and then combining them in the correct order."""
        s = str(paragraph)

        punctuation_indices: dict[int, str] = {
            m.start(1): m.group(1) for m in re.finditer("</abbr>(\W)<abbr", s)
        }  # (\W) matches any non-word character

        # Match anything, even symbols between </abbr> and </p> to get the last character in the paragraph
        end_of_paragraph_indices: dict[int, str] = {
            m.start(1): m.group(1) for m in re.finditer("</abbr>(.*)</p>", s)
        }

        # Find every class="tr">(\w)<... pattern and extract the character
        tr_indices: dict[int, str] = {
            m.start(1): m.group(1)
            for m in re.finditer(r'<span class="tr">(\w*)</span>', s)
        }  # (\w) matches any word character
        # combine the two dictionaries to get the punctuation character and the tr character in the correct order

        combined_idx = punctuation_indices | tr_indices | end_of_paragraph_indices

        return "".join(combined_idx[idx] for idx in sorted(combined_idx.keys()))

    @staticmethod
    def _soup_to_raw_lookup_list(soup: BeautifulSoup) -> list[tuple[str, str, str]]:
        # Each word tag is structured like:
        # China (HSK1)"><ruby><span class="si">中国</span><span class="tr">中國</span><rt>Zhōngguó</rt></ruby></abbr>
        # We want to extract it as (word, pinyin, traditional)
        word_tags = soup.find_all("abbr", {"rel": "tooltip"})

        word_list: list[tuple[str, str, str]] = []
        for t in word_tags:
            translation = t.attrs["title"].split("\n")[-1]

            if hsk_info := re.search(r"HSK(\d+)", translation):
                hsk_level = int(hsk_info[1])

                # Skip easy words
                if hsk_level < MIN_HSK_LEVEL:
                    continue

            traditional_char = t.select_one(".tr").text
            # simplified_char = t.select_one(".si").text

            # pinyin might contain \xa0 (non-breaking space, so replace it with a normal space)
            pinyin = t.select_one("rt").text.replace("\xa0", " ")

            word_list.append((translation, pinyin, traditional_char))

        return word_list

    def _soup_to_word_list(self, soup: BeautifulSoup) -> str:
        lookup_list = self._soup_to_raw_lookup_list(soup)
        word_list = ""
        for word, pinyin, char in lookup_list:
            translation = word.title().replace("Hsk", "HSK")
            # Skip words that start with a parenthesis, because it might be a grammar point
            if translation.startswith("("):
                continue
            word_list += f"{char} ({pinyin}): {translation}\n"
        return word_list


class DocumentWriter:
    def __init__(self):
        # Save each JSON to a docx page
        self.doc = Document()

    def __call__(self, final_texts: list[ArticleTextCollection]) -> None:
        # Define the maximum number of words allowed in each column
        for text in tqdm(final_texts):
            self._add_page_to_doc(text)

        self.doc.save("output.docx")

    def _add_page_to_doc(self, text: ArticleTextCollection) -> None:
        doc = self.doc

        # Assuming you have initialized the 'doc' object
        doc.add_heading(text.chinese_title, level=1)
        subtitle = f"{text.title} - ({text.tags} - {text.hsk_tag})"
        doc.add_heading(subtitle, level=2)

        # Make the main text size 14
        main_text = doc.add_paragraph(text.main_text)
        main_text.style.font.size = Pt(16)

        # Create a new section with two columns
        doc.add_page_break()

        # Set the section type to continuous
        section = doc.add_section(WD_SECTION_START.CONTINUOUS)

        # Define the number of columns and their spacing
        # Ensure the section doesn't start on a new page
        section.start_type.first_page = False
        section.left_margin = Pt(72)  # Set the left margin (adjust as needed)
        section.right_margin = Pt(72)  # Set the right margin (adjust as needed)
        section.cols_number = 2

        # Set the spacing between columns (adjust as needed)
        section.cols_space = Pt(36)

        # Create a table for the word list with two columns
        word_list = doc.add_table(rows=1, cols=2)

        # Access the cells in the first row to fill the word list data
        cells = word_list.rows[0].cells
        words = text.sanitized_word_list

        self._fill_words_into_columns(cells, doc, words)
        # Add a page break after the word list section
        doc.add_page_break()

    @staticmethod
    def _fill_words_into_columns(
        cells: list[Tag], doc: Document, words: set[str]
    ) -> None:
        # Keep track of word count in each column
        word_count_column_1 = 0
        word_count_column_2 = 0
        # Fill the two columns with word list data
        for word in words:
            if word_count_column_1 < MAX_WORDS_PER_COLUMN:
                cells[0].text += word + "\n"
                word_count_column_1 += 1
            elif word_count_column_2 < MAX_WORDS_PER_COLUMN:
                cells[1].text += word + "\n"
                word_count_column_2 += 1
            else:
                # Both columns have reached the maximum number of words, create a new page and table
                doc.add_page_break()
                word_list = doc.add_table(rows=1, cols=2)
                word_list.style.font.size = Pt(10)
                cells = word_list.rows[0].cells
                cells[0].text += word + "\n"
                word_count_column_1 = 1
                word_count_column_2 = 0


class Main:
    def __init__(self) -> None:
        self.fetcher = MainArticleFetcher()
        self.document_writer = DocumentWriter()

    def __call__(self) -> None:
        # First, get the article metadata. That means the title, chinese title, url, tags, and HSK level
        article_metadata = self._get_article_metadata()

        # Then, get the main text and word list for each article
        final_texts = self._get_text_collection(article_metadata)

        self.document_writer(final_texts)

    def _get_text_collection(
        self, article_metadata: list[ArticleMetadata]
    ) -> list[ArticleTextCollection]:
        if FINAL_TEXT_JSON.exists():
            with open(FINAL_TEXT_JSON, "r") as f:
                return [ArticleTextCollection(**x) for x in json.load(f)]

        final_json = []
        for article in tqdm(article_metadata):
            try:
                main_text, word_list = self.fetcher(article.url)
            except Exception:
                continue

            final_json.append(
                ArticleTextCollection(
                    title=article.title,
                    chinese_title=article.chinese_title,
                    url=article.url,
                    tags=article.tags,
                    hsk_tag=article.hsk_tag,
                    main_text=main_text,
                    word_list=word_list,
                )
            )

        FINAL_TEXT_JSON.write_text(
            json.dumps(
                [asdict(x) for x in final_json],
                indent=4,
                ensure_ascii=False,
            )
        )

        return final_json

    def _get_article_metadata(self) -> list[ArticleMetadata]:
        if not METADATA_JSON.exists():
            return self._page_urls_to_json()

        with open(METADATA_JSON, "r") as f:
            metadata_dict = json.load(f)
            return [ArticleMetadata(**article) for article in metadata_dict]

    def _page_urls_to_json(self) -> list[ArticleMetadata]:
        pagenum_param = "&pagenum={}"
        # include only HSK4, HSK5, HSK6 with the psot_tag
        # Get the first page
        first_page = requests.get(BASE_URL + pagenum_param.format(1), headers=HEADERS)
        page_soup = BeautifulSoup(first_page.text, "html.parser")
        max_page = int(
            page_soup.find("div", {"class": "ecs-posts"})
            .attrs["data-settings"]
            .split('"max_num_pages":')[-1]
            .split(",")[0]
        )

        article_metadata: list[ArticleMetadata] = []
        # get the article metadata for all pages
        for page_num in range(1, max_page + 1):
            page = requests.get(
                BASE_URL + pagenum_param.format(page_num), headers=HEADERS
            )
            page_soup = BeautifulSoup(page.text, "html.parser")
            article_metadata += self._get_article_metadata_from_soup(page_soup)

        # Save the article metadata to a JSON file
        metadata_dict = [asdict(article) for article in article_metadata]

        METADATA_JSON.write_text(
            json.dumps(metadata_dict, indent=4, ensure_ascii=False)
        )

        return article_metadata

    @staticmethod
    def _get_article_metadata_from_soup(soup: BeautifulSoup) -> list[ArticleMetadata]:
        article_divs = soup.find_all("article", {"class": "elementor-post"})
        article_metadata = []
        for article in article_divs:
            # <span class="elementor-heading-title elementor-size-default"><a href="https://mandarinbean.com/the-tortoise-and-the-hare/">The Tortoise and the Hare</a></span> </div>
            title = article.find("span", {"class": "elementor-heading-title"}).text

            # <span class="si">龟兔赛跑</span><span class="tr">龜兔賽跑</span></a></span> </div>
            chinese_title = article.find("span", {"class": "tr"}).text

            # find the href url, but only if it is inside <span class="elementor-heading-title elementor-size-default">
            url = (
                article.find("span", {"class": "elementor-heading-title"})
                .find("a")
                .attrs["href"]
            )

            meta_tags = article.attrs["class"]

            all_tags = [
                tag.lstrip("tag-") for tag in meta_tags if tag.startswith("tag-") if tag
            ]

            hsk_tag = next((tag for tag in all_tags if "hsk" in tag)).upper()
            category_tag = next((tag for tag in all_tags if "hsk" not in tag)).title()

            article_metadata.append(
                ArticleMetadata(
                    title=title,
                    chinese_title=chinese_title,
                    url=url,
                    tags=category_tag,
                    hsk_tag=hsk_tag,
                )
            )

        return article_metadata


main = Main()

if __name__ == "__main__":
    main()
